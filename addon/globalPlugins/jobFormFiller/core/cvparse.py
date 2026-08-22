# cvparse.py - pull profile fields out of CV text, for import-and-review.
#
# Deterministic and conservative: it extracts the reliable things (email,
# phone, LinkedIn, a likely name, and rough section blocks) and NOTHING it is
# unsure of. It never fabricates a field. Whatever it returns is shown to the
# user to correct before saving, so the cost of a miss is small, the cost of a
# confident wrong guess is not - hence "extract only what is clear".
#
# Getting text OUT of a .docx or .pdf is a separate step (extract_text below),
# done with bundled libraries in the real add-on. The parsing logic here works
# on plain text so it is fully testable.

import re
import unicodedata

_EMAIL = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_LINKEDIN = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9\-_%]+", re.I)
_PHONE = re.compile(r"(?<!\w)(\+?\d[\d\s().\-]{7,}\d)(?!\w)")

# Section headings per concept per language. Extend a language by adding words.
_HEADINGS = {
    "education": {
        "it": ["istruzione", "formazione"],
        "pt": ["educacao", "formacao", "habilitacoes"],
        "pl": ["wyksztalcenie", "edukacja"],
        "nl": ["opleiding", "onderwijs"],
        "en": ["education", "qualifications", "academic"],
        "fr": ["formation", "education", "diplomes"],
        "de": ["ausbildung", "bildung"],
        "es": ["educacion", "formacion"],
        "ar": ["التعليم", "المؤهلات"],
    },
    "experience": {
        "it": ["esperienza", "esperienza professionale", "esperienza lavorativa"],
        "pt": ["experiencia", "experiencia profissional"],
        "pl": ["doswiadczenie", "doswiadczenie zawodowe"],
        "nl": ["werkervaring", "ervaring"],
        "en": ["experience", "employment", "work history", "work experience"],
        "fr": ["experience", "experience professionnelle", "emploi"],
        "de": ["berufserfahrung", "erfahrung"],
        "es": ["experiencia", "experiencia laboral"],
        "ar": ["الخبرة", "الخبرات"],
    },
    "skills": {
        "it": ["competenze"],
        "pt": ["competencias", "habilidades"],
        "pl": ["umiejetnosci"],
        "nl": ["vaardigheden"],
        "en": ["skills", "technical skills", "competencies"],
        "fr": ["competences"],
        "de": ["kenntnisse", "faehigkeiten"],
        "es": ["habilidades", "competencias"],
        "ar": ["المهارات"],
    },
}


_STROKE = str.maketrans({"ł": "l", "Ł": "l", "ø": "o", "Ø": "o",
                         "đ": "d", "Đ": "d", "ð": "d", "þ": "th"})


def _fold(s: str) -> str:
    s = (s or "").translate(_STROKE)
    s = unicodedata.normalize("NFKD", s)
    return "".join(c for c in s if not unicodedata.combining(c)).lower()


def _heading_key(line: str):
    """Return the section key if this line is a recognised heading, else None."""
    low = _fold(line.strip())
    if not low or len(low) >= 40:
        return None
    for key, langs in _HEADINGS.items():
        for words in langs.values():
            for w in words:
                fw = _fold(w)
                if low == fw or low.startswith(fw):
                    return key
    return None


def _looks_like_name(line: str) -> bool:
    line = line.strip()
    if not line or any(ch.isdigit() for ch in line) or "@" in line:
        return False
    words = line.split()
    if not (2 <= len(words) <= 4):
        return False
    # each word begins with a letter (works for non-latin scripts too)
    return all(w[:1].isalpha() for w in words)


def parse_cv_text(text: str) -> dict:
    result = {}
    lines = [ln.rstrip() for ln in (text or "").splitlines()]

    m = _EMAIL.search(text or "")
    if m:
        result["email"] = m.group(0)

    m = _LINKEDIN.search(text or "")
    if m:
        result["linkedin"] = m.group(0)

    # phone: take the first match that has enough digits to be a real number.
    for m in _PHONE.finditer(text or ""):
        digits = re.sub(r"\D", "", m.group(1))
        if 9 <= len(digits) <= 15:
            result["phone"] = m.group(1).strip()
            break

    # name: first non-empty line that looks like a name.
    for ln in lines:
        if ln.strip():
            if _looks_like_name(ln):
                result["full_name"] = ln.strip()
            break

    # sections: collect lines under a recognised heading until the next heading.
    sections = {}
    current = None
    for ln in lines:
        matched = _heading_key(ln)
        if matched:
            current = matched
            sections.setdefault(current, [])
            continue
        if current and ln.strip():
            sections[current].append(ln.strip())
    for key, body in sections.items():
        if body:
            result[key] = body

    return result


def extract_text(path: str) -> str:
    """Get plain text out of a CV file. .docx and .pdf are what CVs actually
    come in. Both libraries used here are pure Python, so they bundle into the
    add-on cleanly. Image-only (scanned) PDFs yield little or no text and are
    the OCR case, handled by the fallback rung, not here."""
    ext = path.lower().rsplit(".", 1)[-1] if "." in path else ""
    if ext in ("txt", "md"):
        with open(path, encoding="utf-8", errors="replace") as f:
            return f.read()
    if ext == "docx":
        import docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:                    # CVs often use tables
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    raise NotImplementedError(f"unsupported CV format: .{ext}")
